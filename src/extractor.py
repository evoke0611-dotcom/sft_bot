from pathlib import Path
from typing import List, Dict, Any

import pandas as pd
from pypdf import PdfReader
from docx import Document

from src.cleaner import clean_text
from src.config import SUPPORTED_EXTENSIONS


def extract_pdf(file_path: Path) -> List[Dict[str, Any]]:
    """
    Extract text page-wise from PDF.
    """

    documents = []
    reader = PdfReader(str(file_path))

    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = clean_text(text)

        if text:
            documents.append({
                "text": text,
                "metadata": {
                    "source_file": file_path.name,
                    "file_path": str(file_path),
                    "file_type": ".pdf",
                    "page": page_number
                }
            })

    return documents


def extract_docx(file_path: Path) -> List[Dict[str, Any]]:
    """
    Extract paragraphs and tables from DOCX.
    """

    doc = Document(str(file_path))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                parts.append(" | ".join(row_text))

    text = clean_text("\n".join(parts))

    if not text:
        return []

    return [{
        "text": text,
        "metadata": {
            "source_file": file_path.name,
            "file_path": str(file_path),
            "file_type": ".docx",
            "page": None
        }
    }]


def extract_file(file_path: Path) -> List[Dict[str, Any]]:
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return extract_pdf(file_path)

    if suffix == ".docx":
        return extract_docx(file_path)

    return []


def load_all_documents(raw_data_dir: Path) -> List[Dict[str, Any]]:
    """
    Loads all supported files from data/raw recursively.
    """

    all_documents = []

    for file_path in raw_data_dir.rglob("*"):
        if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
            print(f"Extracting: {file_path.name}")
            docs = extract_file(file_path)
            all_documents.extend(docs)

    return all_documents